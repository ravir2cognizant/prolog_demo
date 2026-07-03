"""
build-design-reference.py

Generates three project-independent design reference .docx files for the agentic
delivery pipeline, distilled from .claude/kb/agentic-delivery-core-kb.md and
.claude/kb/cost-optimization-kb.md.

Outputs (written to poc-workspace root):
  1. Agentic-Pipeline-Design-Handbook.docx        ~50 pages, comprehensive
  2. Agentic-Pipeline-Concise-Playbook.docx       ~20 pages, focused
  3. Agentic-Pipeline-Cheat-Sheet.docx            ~3 pages, at-the-keyboard quick ref

Run with:  py -3.14 agentic-pipeline/scripts/build-design-reference.py
(The free-threading variant `py -3.14t` lacks a working lxml etree; use plain 3.14.)
"""

from __future__ import annotations

import os
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Styling helpers
# ---------------------------------------------------------------------------

VERSION = "1.0"
TODAY = date.today().isoformat()

COLOR_PRIMARY = RGBColor(0x1F, 0x3A, 0x5F)   # deep blue for headings
COLOR_ACCENT = RGBColor(0x2E, 0x7D, 0x32)    # green for callouts
COLOR_WARN = RGBColor(0xB7, 0x1C, 0x1C)      # red for anti-patterns
COLOR_MUTED = RGBColor(0x55, 0x55, 0x55)     # grey for captions


def _set_cell_shading(cell, hex_color: str) -> None:
    """Apply a background shade (e.g. '1F3A5F') to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _configure_base_styles(doc: Document) -> None:
    """Tighten default fonts and heading colors."""
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    for lvl, size, bold in [
        ('Title', 28, True),
        ('Heading 1', 18, True),
        ('Heading 2', 14, True),
        ('Heading 3', 12, True),
        ('Heading 4', 11, True),
    ]:
        s = doc.styles[lvl]
        s.font.name = 'Calibri'
        s.font.size = Pt(size)
        s.font.bold = bold
        if lvl != 'Title':
            s.font.color.rgb = COLOR_PRIMARY


def add_title_block(doc: Document, title: str, subtitle: str, audience: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = COLOR_PRIMARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.italic = True
    r.font.size = Pt(14)
    r.font.color.rgb = COLOR_MUTED

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Version {VERSION}   ·   {TODAY}").italic = True

    doc.add_paragraph()
    aud = doc.add_paragraph()
    aud.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = aud.add_run("Audience: ")
    r.bold = True
    aud.add_run(audience)

    doc.add_paragraph()
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run(
        "This document is project-independent. It captures the philosophy, "
        "principles, decisions, and patterns of a multi-agent software delivery "
        "pipeline so the same playbook can be applied to future projects on any "
        "tech stack."
    )
    r.italic = True
    r.font.color.rgb = COLOR_MUTED


def add_h1(doc, text):
    doc.add_heading(text, level=1)


def add_h2(doc, text):
    doc.add_heading(text, level=2)


def add_h3(doc, text):
    doc.add_heading(text, level=3)


def add_para(doc, text, italic=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.add_run(item)


def add_callout(doc, title, body, color_hex='1F3A5F'):
    """Single-cell shaded box for emphasis."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    cell = tbl.cell(0, 0)
    _set_cell_shading(cell, 'EAF1F8')
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = COLOR_PRIMARY
    cell.add_paragraph(body)
    doc.add_paragraph()


def add_table_with_header(doc, headers, rows, col_widths_in=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_shading(hdr[i], '1F3A5F')

    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            tbl.rows[ri].cells[ci].text = str(val)

    if col_widths_in:
        for col_idx, w in enumerate(col_widths_in):
            for r in tbl.rows:
                r.cells[col_idx].width = Inches(w)

    doc.add_paragraph()
    return tbl


def add_decision(doc, did, title, rationale, how_to_apply=None):
    """Render one decision row in a stable format."""
    p = doc.add_paragraph()
    r = p.add_run(f"{did} — {title}")
    r.bold = True
    r.font.color.rgb = COLOR_PRIMARY

    p = doc.add_paragraph()
    r = p.add_run("Why: ")
    r.bold = True
    p.add_run(rationale)

    if how_to_apply:
        p = doc.add_paragraph()
        r = p.add_run("How to apply: ")
        r.bold = True
        p.add_run(how_to_apply)
    doc.add_paragraph()


def add_page_break(doc):
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Content blocks (project-independent prose)
# ---------------------------------------------------------------------------

PHILOSOPHY_OPENING = (
    "Agentic delivery is the practice of decomposing a software delivery pipeline into a "
    "network of specialist AI agents, each with a single clearly defined responsibility, "
    "coordinated by a central Orchestrator. The goal is not to automate everything — it is "
    "to make every step explicit, auditable, and reliably reproducible. The pipeline trades "
    "the illusion of one omniscient assistant for a discipline of small, focused, replaceable "
    "specialists whose handoffs are visible on disk."
)

FOUR_PILLARS = [
    ("Single Responsibility",
     "Every agent does one thing and does it well. An agent that does two things does neither "
     "as well as two dedicated agents would. Single-responsibility agents are predictable, "
     "replaceable, independently improvable, and easy to reason about when something goes wrong. "
     "If an agent's stated responsibility contains the word 'and', it is two responsibilities and "
     "must be split."),
    ("Centralised Coordination",
     "One Orchestrator holds all state. It knows every path, every dependency, every task status, "
     "every clarification in flight, every blocker raised. Agents do not know about each other. "
     "They know only their job, their input, and their output — and they ask the Orchestrator for "
     "all of that. Centralised state makes debugging trivial and makes the pipeline easy to modify."),
    ("Explicit over Implicit",
     "Nothing in a well-designed agentic pipeline is assumed, guessed, or inferred silently. "
     "Every ambiguity is raised as a clarification. Every completion is validated against a "
     "Definition of Done. Every decision is logged in an audit trail. Implicit behaviour is the "
     "enemy of reliable delivery."),
    ("Cost Discipline",
     "Every token spent on coordination is a token not spent on delivery. The pipeline defaults to "
     "the cheapest activation mechanism that achieves the outcome — foreground mode-switch over "
     "sub-agent spawn, hash-skip over re-execution, persisted briefings over re-derivation. Cost "
     "discipline is the difference between a pipeline that runs sustainably and one that bankrupts "
     "the budget after three sprints."),
]

GUIDING_PRINCIPLES = [
    "An agent must never contact another agent directly. All communication routes through the Orchestrator.",
    "An agent must never guess or proceed on an assumption. Uncertainty is always raised explicitly.",
    "An agent must never overwrite output when input has not changed. Idempotency is a first-class property.",
    "An agent must never report complete without passing its Definition of Done checklist.",
    "The Orchestrator must never produce business output. Its single responsibility is coordination.",
    "Paths, dependencies, and context are provided by the Orchestrator — never hardcoded in agent prompts.",
    "Mechanical work (schema validation, hash computation, folder creation, alignment checks) belongs in hook scripts, not in agent prompts. Agents focus on semantic judgement.",
    "Every significant decision is documented with its rationale. Undocumented decisions become mysteries.",
    "Human involvement is minimised but never eliminated. The human is the final escalation path.",
    "The pipeline never starts autonomously. Every sprint begins with a deliberate human signal.",
    "State lives in files on disk, not in agent memory. Agent sessions are stateless; the file system is the durable record.",
    "Sub-agent spawn is the exception, never the default. Foreground mode-switch is the cheapest activation mechanism.",
]

WHEN_TO_USE = [
    ("Use agentic delivery when…",
     [
         "The delivery pipeline has clearly separable phases with defined inputs and outputs between them.",
         "Quality gates between phases are important — a misunderstood requirement is expensive to fix later.",
         "The team wants an auditable, reproducible process new members can follow without tribal knowledge.",
         "Multiple specialists (design, backend, frontend, test, review) work on the same requirements.",
         "Cost is a real constraint and you need to control LLM spend deliberately.",
     ]),
    ("Do not use agentic delivery when…",
     [
         "The work is exploratory or research-oriented with no clear input/output structure.",
         "The team is one person doing everything — coordination overhead exceeds the benefit.",
         "Requirements change so rapidly a structured pipeline cannot keep up.",
         "There is no patience for the upfront investment in agent definitions, hooks, and validators.",
     ]),
]

REFERENCE_AGENTS = [
    ("A-00", "Orchestrator", "Coordinate everything. Hold all state. Produce no business output.", "—"),
    ("A-SM", "Sprint Manager", "Detect sprint start, register sprint, produce velocity report.", "—"),
    ("A-01", "Requirement Analyst", "Read raw inputs, produce structured requirement cards.", "No (split)"),
    ("A-01r", "Requirement Resolver", "Resolve routed clarifications + concerns; emit verdicts.", "No"),
    ("A-02", "API/BFF Designer", "Design API endpoint contracts from requirement cards.", "Yes"),
    ("A-03a", "UI Style Compiler", "Compile design tokens + theme + style-system docs.", "No"),
    ("A-03b", "UI Component Inventory", "Produce component inventory from RC + style-system.", "Yes"),
    ("A-04", "Frontend Developer", "Implement the UI from RC + CI + style-system.", "Yes"),
    ("A-05", "Backend Developer", "Implement the API layer from RC + ED.", "Yes"),
    ("A-06", "Code Reviewer", "Review implementation, produce findings + JSON routing summary.", "No"),
    ("A-07", "Frontend Tester", "Plan + execute FE tests; emit defects + JSON routing summary.", "Yes"),
    ("A-08", "API/BFF Tester", "Plan + execute API tests; emit defects + JSON routing summary.", "Yes"),
]

AGENT_ANATOMY = [
    ("Definition file",
     "States the agent's single responsibility, position in the pipeline, input contract, output contract, "
     "escalation chain, and which protocols apply. Narrative — answers WHAT and WHY."),
    ("Skills file",
     "Domain knowledge, output format specification, quality standards, Definition of Done checklist, "
     "worked examples. Operational — answers HOW."),
    ("Hook script (H-XX)",
     "Deterministic PowerShell (or shell) script that verifies dependencies, resolves paths, computes "
     "the input hash, manages partial output recovery, and returns PROCEED / NO_CHANGE / BLOCKED / "
     "ALIGNMENT_CONFLICT. Optional -PostCheck switch runs declared validators after completion."),
    ("Activation file (CLAUDE-A-XX)",
     "The entry-point file the human reads aloud as 'Activate <agent name>'. Names the definition + "
     "skills + briefing path, declares the default model tier."),
    ("Persisted briefing",
     "A per-task file (T-###-A-##-…-briefing.md) the Orchestrator writes before activation. The agent "
     "reads the briefing once and treats it as authoritative — it never re-derives context."),
    ("Validator(s)",
     "Per-producer V-XX-<topic>.ps1 and cross-cutting V-shared-<topic>.ps1 scripts that mechanically "
     "check schemas and joint contracts. Invoked by the hook pre-activation (alignment) and "
     "post-completion (schema)."),
]

THREE_FOLDER = [
    ("agentic-pipeline/",
     "Pipeline infrastructure: orchestrator manifest, audit log, agent definitions and skills, hook "
     "scripts, validators, helper scripts, briefings, NOTIFICATIONS.md. Versioned with the pipeline, "
     "not the application. THIS IS THE ONLY FOLDER A FRESH WORKSPACE HAS — the other two appear "
     "lazily as a side-effect of the first activation that needs them."),
    ("sprints/sprint-##/  (lazy)",
     "All sprint-scoped artefacts: requirement cards, designs, component inventories, test cases, "
     "test results, defects, disputes, review findings, velocity reports. Isolated per sprint. "
     "Sprint N+1 never overwrites Sprint N. The sprints/ root is created by start-sprint.ps1 on the "
     "first sprint; each sprint-## subfolder is created by start-sprint.ps1 on that sprint's start."),
    ("app/  (lazy)",
     "Long-lived application code. Accumulates across sprints. Owned by developer agents. "
     "Hash files in this folder are sprint-scoped (.input-hash-sprint-##) so cross-sprint runs do "
     "not invalidate each other. The app/ root is NOT created at workspace bootstrap; "
     "app/frontend/ is created by the frontend developer's hook on first activation, and "
     "app/backend/ by the backend developer's hook on first activation."),
]

LAZY_CREATION_RULE = (
    "Lazy creation — a brand-new workspace has only agentic-pipeline/ at the root. The app/ and "
    "sprints/ folders are NOT pre-created by workspace bootstrap. They appear as a side-effect of "
    "the first activation that needs them: sprints/ on the first sprint start, app/frontend/ on the "
    "first frontend developer activation, app/backend/ on the first backend developer activation. "
    "This rule keeps the 'infrastructure only' state visually unambiguous, prevents accidental "
    "commits of empty placeholder directories, and lets a new project start from a clean "
    "infrastructure-only template with no manual cleanup."
)

PROTOCOLS = [
    ("Protocol 1 — Startup",
     "Every agent's first action is to read the orchestrator-supplied briefing (input path, output path, "
     "dependency status, resolved clarifications since last run) and confirm the four core questions: "
     "Where is my input? Where do I write my output? Are my dependencies complete? What has changed since "
     "I last ran? Hook returns are honoured — NO_CHANGE means exit immediately."),
    ("Protocol 2 — Sign-off",
     "Signing agents review the Requirement Analyst's output (RC cards) in READ-ONLY mode at the sign-off "
     "gate. Each consuming downstream agent confirms 'I can do my job with this requirement as written.' "
     "Misunderstood requirements caught at the gate cost 5–10× less than at implementation."),
    ("Protocol 3 — Clarification",
     "When an agent encounters uncertainty, it raises a Clarification (CL-###) via the Orchestrator. The "
     "Orchestrator routes per the agent's escalation chain. Never guess. Never proceed on assumption. "
     "Concerns (CNC-###) are a distinct artefact for 'the source material has a gap' rather than 'I need "
     "an answer'."),
    ("Protocol 4 — Completion",
     "An agent reports complete only after self-validating its Definition of Done. Hook-driven "
     "post-completion validators then mechanically check schemas before the Orchestrator marks the task "
     "complete. False completions (red gates, missing artefacts, schema violations) are caught here."),
    ("Protocol 5 — Cost Discipline",
     "Operate at the lowest cost tier that achieves the outcome. Foreground mode-switch is the default; "
     "sub-agent spawn is the exception. Hash-skip on identical inputs. Persisted briefings. /compact "
     "before context bloat. Sub-agent budget per sprint is enforced and audited."),
]

COST_TIERS = [
    ("T0", "Hash-skip (NO_CHANGE) — hook exits, no LLM run", "~0"),
    ("T1", "Foreground mode-switch — same session, new role", "1× baseline"),
    ("T2", "Foreground with /compact — recover window", "1.1×"),
    ("T3", "Fresh session, state reloaded from disk", "1.2×"),
    ("T4", "Sub-agent spawn (justified)", "3–5×"),
    ("T5", "Sub-agent with inline-return doubling", "5–8×"),
    ("T6", "Sub-agent re-spawn after truncation", "8–12×"),
]

SPAWN_CASES = [
    ("Case A — True parallelism",
     "Two or more agents must run simultaneously because their work is independent and sequencing them "
     "would idle one for hours. Canonical example: parallel rework after consolidated review + test feedback."),
    ("Case B — Heavy context isolation",
     "Output exploration would blow up the foreground context. Canonical example: an Explore agent "
     "scanning 100+ files where returning all matches inline would exceed the foreground working budget. "
     "Specify exploration breadth so the sub-agent right-sizes its context."),
    ("Case C — Truncation-risk payload",
     "A single inline emission would exceed the transcript ceiling and cause silent tail truncation. "
     "Split into 2–3 focused parallel sub-agents AND emit ledger-first AND prefer direct-disk-write."),
]

SPAWN_ANTIPATTERNS = [
    "\"It feels cleaner to delegate this.\"",
    "\"The foreground agent is busy.\" (It isn't — it's the same session.)",
    "\"The sub-agent has a tuned prompt for this.\" (Mode-switch achieves the same.)",
    "\"I want to keep the main context tidy.\" (Use /compact instead.)",
    "\"This task has 2–3 tool calls and I'd rather isolate it.\" (Foreground tool calls are cheaper.)",
]

COORDINATION_PATTERNS = [
    ("Requirement Sign-off Gate",
     "After the Requirement Analyst emits RC cards, ALL signing agents review them READ-ONLY in parallel. "
     "Gate opens only when every signing agent has signed every card (or partial gate per fully-signed "
     "card). Catches ambiguity before any design begins."),
    ("Clarification Loop",
     "Agent raises CL → Orchestrator routes per escalation chain → resolver returns verdict "
     "(RESOLVED_FROM_SOURCE | NEEDS_RC_UPDATE | HUMAN_BLOCKER) → Orchestrator forwards to requester or "
     "raises HB. The agent never contacts another agent directly."),
    ("Consolidated Rework Loop",
     "Code-review findings AND test defects feed into a single rework cycle. Orchestrator reads "
     "review-summary.json + defect-summary-<layer>.json and emits a unified routing plan. Iterative "
     "cycles (T-007b, T-007c) are permitted and visible in IDs."),
    ("Defect Dispute Mechanism",
     "When a developer disagrees with a test-raised defect, they author a DSP-*.md with a verdict "
     "(not-a-defect | test-case-incorrect | requirement-mismatch | valid-defect). The originating test "
     "agent re-judges. requirement-mismatch escalates to the resolver. Audit trail captures the contest."),
    ("Parallel Activation",
     "The Orchestrator activates multiple agents simultaneously when they have no dependency on each "
     "other. The manifest dependency graph is authoritative. Cost note: parallel via sub-agent spawn "
     "counts as ONE Case-A spawn under the budget, not two."),
    ("Test Sign-off Currency",
     "When a test agent reports PASS, the hook commits a .signoff-hash capturing app + test-case state. "
     "After every developer-agent completion, the test hook re-evaluates: NO_CHANGE means sign-off is "
     "still current; PROCEED means re-test. Zero-LLM regression detection."),
]

VALIDATION_LAYERS = [
    ("Pre-activation (alignment)",
     "Hook invokes V-shared-*-alignment.ps1 before returning PROCEED. Misalignment returns "
     "ALIGNMENT_CONFLICT; the consuming agent never starts on misaligned inputs. The Orchestrator "
     "routes the conflict back to the producing agents."),
    ("Post-completion (schema)",
     "Producer hook accepts -PostCheck. The Orchestrator invokes it after the agent reports complete, "
     "BEFORE marking the task [x]. The hook runs the declared Tier-1 validators and emits a single "
     "VALIDATION_PASS / VALIDATION_FAIL signal."),
    ("Self-validation (Definition of Done)",
     "Inside the agent. The skills file's DoD checklist is the contract. Self-validation is the "
     "agent's responsibility but is NEVER the only check — hooks always re-verify mechanically."),
]

FAILURE_MODES = [
    ("Agent guessing instead of clarifying",
     "Symptom: output cites assumptions or 'I'll assume X'. Cause: vague prompts or missing escalation chain.",
     "Prevention: every agent definition has an explicit clarification protocol and escalation chain. "
     "Detection: review responses for 'assume' / 'inferred' language.",
     "Recovery: roll back; require the agent to raise CL-### before retrying."),
    ("False completion reports",
     "Symptom: agent reports complete but Definition of Done is not met. Cause: skipping the self-validation step.",
     "Prevention: hook-driven post-completion validators run regardless of the agent's claim. "
     "Detection: VALIDATION_FAIL on the -PostCheck invocation.",
     "Recovery: mark task [V] Validation Failed; re-activate the agent with the failure detail in its next briefing."),
    ("Hash file corruption",
     "Symptom: hook reports PROCEED on inputs that haven't changed. Cause: hash file deleted, partial-write interrupted, or hash algorithm mismatch.",
     "Prevention: hooks write hash atomically and clean partial outputs when hash file is missing.",
     "Recovery: delete the suspect hash, re-run; the agent will regenerate and rewrite the hash."),
    ("Pipeline drift across sessions",
     "Symptom: a re-spawned agent contradicts a previous decision. Cause: not reading the persisted briefing.",
     "Prevention: every activation reads the briefing first; the briefing names prior decisions verbatim.",
     "Recovery: rebuild the briefing from the manifest + audit-log; re-spawn."),
    ("Sign-off gate bypassed",
     "Symptom: design started before all sign-offs collected. Cause: time pressure.",
     "Prevention: the hook explicitly checks T-GATE = [x] before returning PROCEED.",
     "Recovery: pause downstream; complete sign-off; the partial design counts as a CR for rework."),
    ("Manifest state corruption",
     "Symptom: a task shows two states or a section ends mid-row. Cause: concurrent writes (non-Orchestrator agent wrote).",
     "Prevention: single-writer rule — only A-00 writes the manifest. A-SM has a narrow carve-out for Sprint Registry + ACTIVE SPRINT.",
     "Recovery: restore from the most recent clean audit-log marker."),
    ("Sub-agent write denial",
     "Symptom: spawned sub-agent's writes silently fail. Cause: write/edit permissions not pre-authorised.",
     "Prevention: pre-auth Write/Edit in settings.local.json BEFORE launching any sub-agent.",
     "Recovery: re-spawn after fixing permissions; verify writes landed on disk."),
    ("Sub-agent context blow-up",
     "Symptom: sub-agent returns truncated or partial output. Cause: scope too wide for one spawn.",
     "Prevention: split into 2–3 focused sub-agents (Recovery-A scaffold, Recovery-B features, Recovery-C tests). Emit ledger-first.",
     "Recovery: re-run only the truncated slice; ledger-first emission lets you identify what survived."),
    ("Cross-agent contract drift in parallel rework",
     "Symptom: A-04 and A-05 disagree on a shape after T-007. Cause: no canonical source-of-truth was pre-decided.",
     "Prevention: Orchestrator pre-decides canonical side (server-canonical default for response shapes) and embeds in BOTH briefings verbatim.",
     "Recovery: re-route the disputed shape with the canonical side stated; the non-canonical agent does only the read-side adaptation."),
    ("Dispute deadlock (test vs developer)",
     "Symptom: a DSP-*.md bounces past two rounds without resolution. Cause: ambiguous requirement.",
     "Prevention: the test agent's first response to a dispute cites the specific requirement text it relied on.",
     "Recovery: escalate to the requirement resolver via the Orchestrator; if NEEDS_RC_UPDATE or HUMAN_BLOCKER, the requirement itself is the root cause."),
    ("Model mis-tier",
     "Symptom: a sprint costs more than projected, OR quality drops on a stable task. Cause: declared model tier wrong, or dynamic rule mis-firing.",
     "Prevention: every sub-agent spawn writes a model-selection audit row with reason.",
     "Recovery: adjust the declared tier; review the rework-cycle counter if the dynamic rule is firing wrongly."),
    ("Validator-fail loop",
     "Symptom: VALIDATION_FAIL on the same artefact across multiple agent retries. Cause: the agent doesn't understand the violation, OR the validator itself has a bug.",
     "Prevention: validators emit specific error codes + targets + details; briefings include the failure verbatim.",
     "Recovery: if 3 rounds fail with the same error, treat the validator as suspect; human-review the validator logic."),
]

ANTIPATTERNS = [
    "Combined agents — merging two specialist responsibilities into one to save 'overhead'.",
    "Hardcoded paths — filesystem paths in agent prompts or skill files.",
    "Guessing — accepting 'I'll assume X' from an agent without raising a clarification.",
    "Bypassed gates — skipping the sign-off gate under time pressure.",
    "Autonomous sprint starts — allowing the pipeline to start without a human trigger.",
    "Shared manifest writes — any agent other than the Orchestrator (or A-SM in its carve-out) writing the manifest.",
    "Stale context — activating an agent without preparing and persisting a context briefing.",
    "Silent gaps — shipping a card that misses a stated constraint without a documented exception.",
    "Sub-agent-only writes — relying on a background sub-agent to persist artefacts without verifying write propagation.",
    "Markdown content parsing for routing — parsing review/defect Markdown bodies in coordination logic instead of JSON summaries.",
    "Mechanical checks in agent prompts — burning LLM tokens on schema validation.",
    "Rule sprawl in model-selection — adding a second dynamic model-tier rule without explicit justification.",
    "Side-car patch files left un-merged — *-<topic>.md patch files intended for a parent file but never merged.",
    "Spawning for sequential dependent work — pure duplication with no parallelism benefit.",
    "Sub-agent for tasks under 3 tool calls — spawn overhead exceeds the work itself.",
    "Re-loading KBs in the same session — KBs are already in context; re-reading is waste.",
    "Re-running an agent because briefing felt thin — strengthen the briefing, not the run count.",
    "Mode-switching 6+ times without /compact — context bloat costs more than a fresh session.",
    "Multiple HBs for decisions that emerged together — use a composite HB.",
]

# Decision log — every decision with rationale + practical guidance
DECISIONS = [
    # v1.0 — Foundational
    ("D-001", "Orchestrator is the single source of truth",
     "Centralised state makes debugging trivial. Distributed manifest writes produce race conditions and untraceable history.",
     "Only the Orchestrator (and Sprint Manager within its narrow carve-out) writes the manifest. Every other agent emits signals; the Orchestrator appends."),
    ("D-002", "Single responsibility per agent",
     "Multi-responsibility agents produce inconsistent quality. A prompt that asks for 'design and implementation' optimises for neither.",
     "If the agent's stated responsibility contains 'and', split it. Test against the three SRP questions before creating an agent."),
    ("D-003", "Agents never contact each other directly",
     "All communication routes through the Orchestrator for audit + logging. Direct agent-to-agent calls create hidden state.",
     "Every clarification, every handoff, every blocker goes through the Orchestrator. Agent definitions name escalation chains, not destinations."),
    ("D-004", "Requirement sign-off gate before any design work",
     "Misunderstood requirements caught at design cost 5–10× less than at implementation. Surfacing ambiguity early is the highest-leverage check in the pipeline.",
     "Every downstream consumer reviews requirement cards READ-ONLY at the gate. The gate opens only when every signing agent signs every card."),
    ("D-005", "Hooks scripts for dependency and path resolution",
     "Move environment resolution out of the agent prompt into a deterministic script. Hardcoded paths in prompts rot the first time a folder moves.",
     "Every agent has an H-XX-<fullname>.ps1 that verifies dependencies, resolves paths, computes hashes, and returns PROCEED / NO_CHANGE / BLOCKED."),
    ("D-006", "Input hashing to prevent unnecessary re-runs",
     "Identical inputs produce identical outputs. Re-running on unchanged inputs is pure waste — token cost without value.",
     "Hooks compute MD5 of all inputs, compare to .input-hash, return NO_CHANGE if matched. Agents trust NO_CHANGE absolutely — they do not 'double-check'."),
    ("D-007", "Sprint-scoped folder structure",
     "Prevents Sprint N+1 from accidentally overwriting Sprint N outputs. Sprint isolation is a hard guarantee, not a convention.",
     "All sprint artefacts go under sprints/sprint-##/. Long-lived code under app/ uses sprint-scoped hash files (e.g. .input-hash-sprint-##)."),
    ("D-008", "START_SPRINT file as explicit sprint trigger",
     "Human control over sprint scope is preserved. The pipeline never starts a new sprint on its own.",
     "Sprint Manager activates only when a START_SPRINT file appears in the sprint-##/inputs/ folder. The human creates that file deliberately."),
    ("D-009", "Clarification escalation chain per agent",
     "Routing every clarification directly to the Requirement Analyst creates a bottleneck. Most clarifications can be resolved by an upstream peer.",
     "Each agent definition names a primary escalation target (usually the immediate upstream producer). Only unresolved escalations reach the resolver, then human."),
    ("D-010", "Context briefing at every agent activation",
     "Agent sessions are stateless. Briefings prevent re-raising already-resolved questions and lose-the-thread restarts.",
     "Before activation, the Orchestrator persists a T-###-A-##-<fullname>-briefing.md containing inputs, dependencies, resolved CLs, RC versions, and prior-attempt history."),
    ("D-011", "Three-folder workspace separation",
     "Different characteristics, different lifecycles, different ownership. Long-lived code, sprint-scoped artefacts, and pipeline infrastructure mix badly.",
     "app/ — code, accumulates across sprints. sprints/sprint-##/ — sprint-scoped artefacts. agentic-pipeline/ — pipeline infrastructure."),
    # v1.2 — Briefings, carve-outs, concern artefact
    ("D-012", "Briefings persisted as files",
     "Enables pause/resume across sessions. Gives every activation a recoverable record. Lets the human debug an agent's behaviour by reading what it was told.",
     "Briefings live in agentic-pipeline/briefings/T-###-A-##-<fullname>-briefing.md. The file is the durable record; the prompt-embedded copy is the working copy."),
    ("D-013", "Sprint Manager may write Sprint Registry + ACTIVE SPRINT",
     "Sprint Manager owns the sprint lifecycle. A narrow, explicit carve-out from the single-writer rule keeps lifecycle changes coherent.",
     "A-SM is the only non-Orchestrator agent allowed to write specific manifest sections. Everything else still routes through A-00."),
    ("D-014", "Concern (CNC-###) is a distinct artefact from Clarification (CL-###)",
     "Clarification is 'I need an answer'. Concern is 'the source material has a gap'. They resolve differently and have different audiences.",
     "Use CL when a peer agent can resolve. Use CNC when the source artefact itself needs revision."),
    ("D-015", "Composite human blockers are permitted",
     "Resolving related decisions in batches avoids gate churn. Five HBs raised one-at-a-time on the same topic burn human attention.",
     "Bundle related decisions into one HB. Resolve once. Cascade RC version bumps in a single batch."),
    ("D-016", "Non-blocking human blockers are permitted",
     "Some ambiguities deserve human attention without halting work. A 'while you're here, please confirm X' note is cheaper than a full pause.",
     "Mark the HB non-blocking. The pipeline continues; the human resolves at next checkpoint."),
    ("D-017", "Documented compliance exceptions are first-class",
     "A visible, documented exception beats a hidden gap. The audit trail must show every constraint that wasn't fully met and why.",
     "Record exceptions in the RC card body and the gate's Cross-card notes. The DoD treats a documented exception as a passing condition for the affected rule."),
    ("D-018", "Foreground-write pattern for sub-agents (with direct-disk-write update)",
     "Default to foreground-write because it preserves the parent's view of what was written. Direct-disk-write only when pre-auth is verified and the payload is large.",
     "Pre-authorise Write/Edit in settings.local.json BEFORE launching any sub-agent. Apply direct-disk-write for >50KB payloads."),
    # v1.3 — Rework, model tier, truncation
    ("D-019", "Canonical source-of-truth decision before parallel rework",
     "Without a pre-decision, parallel agents drift on shared findings — one fixes the server side, the other fixes the client side, and they end up inconsistent.",
     "Activating agent pre-decides canonical side for every 'shared' finding (default: server for response shapes; producer for protocol/middleware; component-owner for UI contracts). Embed the decision verbatim in BOTH briefings."),
    ("D-020", "Rework agents emit ledger JSON + Excel report",
     "Single canonical record. Ledger emitted FIRST in the response so it survives transcript truncation. Excel for humans, JSON for the Orchestrator.",
     "Every rework agent writes a *-ledger.json AND a *-rework-report.xlsx at known paths. The ledger is the routing contract."),
    ("D-021", "Iterative rework (T-007b, T-007c, …) is permitted",
     "Real rework cycles introduce new findings. Pretending one round is enough produces zombie defects that ship.",
     "The iteration count is visible in task IDs. Suffix the task ID (T-007 → T-007b → T-007c) and the artefact prefix (CR-* → CR2-* → CR3-*)."),
    ("D-022", "Verification gates (lint + tests) are preconditions to rework completion",
     "A rework agent reporting complete with red gates is producing a false completion. The pipeline must catch this.",
     "Definition of Done checklist includes lint clean + tests green. Hook post-completion can be augmented to assert these."),
    # v2.0 — 12-agent pipeline, validators, model tier, naming
    ("D-023", "Protocol 5 (Cost Discipline) is mandatory",
     "Without cost discipline, multi-agent pipelines spend 4–9× their necessary cost. Every coordination decision has a cost tier.",
     "Foreground mode-switch is default. Sub-agent spawn is exception (Case A/B/C). Hash-skip absolutely. Persisted briefings. /compact proactively. Full rules in cost-optimization-kb.md."),
    ("D-024", "Mechanical schema checks belong in hook validators, not agent prompts",
     "Validators are deterministic and cost zero LLM tokens. Agents wasting tokens re-checking frontmatter is the textbook misuse.",
     "Per-producer V-<id>-<topic>.ps1 and cross-cutting V-shared-<topic>.ps1 scripts. Hooks invoke them pre-activation (alignment) and post-completion (schema)."),
    ("D-025", "Producer JSON routing summaries",
     "The Orchestrator parsing Markdown bodies for routing decisions is fragile and slow. JSON summaries are a stable contract.",
     "Every producer emits a JSON summary alongside Markdown artefacts (review-summary.json, defect-summary-<layer>.json, dispute-summary.json, cross-sprint-refs.json). Orchestrator reads ONLY JSON for routing."),
    ("D-026", "Split UI agent into Style Compiler + Component Inventory",
     "Style tokens and component inventories change at different rates. Independent hashes let styles re-compile without invalidating components and vice versa.",
     "A-03a (UI Style Compiler) reads only style inputs. A-03b (UI Component Inventory) reads RC + A-03a outputs. Each has its own hash."),
    ("D-027", "Split Requirement Analyst into producer + resolver",
     "Producer is one of the heaviest agents. Resolver has small focused context. Forcing the producer to handle resolutions wastes context.",
     "A-01 produces RC cards. A-01r resolves routed clarifications and concerns. A-01r is cheap (haiku-tier-suitable)."),
    ("D-028", "Test agents are dual-phase (planning + execution) and use a defect dispute mechanism",
     "Test plans should run alongside design (review the RC for testability before code exists). Execution runs after implementation. Disputes prevent dev-vs-test silent disagreement.",
     "Phase 1: emit TC-*.md test cases. Phase 2: execute, emit TR-*.md results + DEF-*.md defects + defect-summary-<layer>.json. Disputes via DSP-*.md with verdicts."),
    ("D-029", "T-007 rework is consolidated (code-review + test defects in one pass)",
     "Running code-review rework and test-defect rework as two separate cycles doubles the rework cost and confuses canonicality.",
     "Orchestrator reads review-summary.json + defect-summary-*.json after both upstream tasks complete, emits a unified T-007 routing plan, activates affected developer(s) once."),
    ("D-030", "Per-agent model tier declaration + one dynamic rework rule",
     "Most agents do not need top-tier models. Reserving Opus for vision-heavy producers and rework recovery achieves quality at ~35–55% of all-Opus cost.",
     "Each activation file declares 'model: haiku | sonnet | opus'. Sub-agent spawns consult a select-model script. One dynamic rule: developer agents in rework cycle ≥ 2 force Opus. Adding a second dynamic rule requires an ADR."),
    ("D-031", "Pipeline filename naming convention with prefix + fullname",
     "Pipeline infrastructure files benefit from a sortable, scannable convention. Agent IDs, hook IDs, validator IDs, and task IDs should all be unambiguous in one glance.",
     "Pipeline infrastructure uses A-/H-/V-/T- prefixes with descriptive fullnames. Sprint artefacts (RC, ED, CI, CR, DEF, etc.) keep their content-type prefixes."),
    ("D-032", "NOTIFICATIONS.md is a single-writer file",
     "Concurrent writes to the human-facing notifications file produce racing or overwritten messages. The notification stream must be coherent.",
     "Orchestrator is the sole writer of NOTIFICATIONS.md. Other agents emit signals (via completion reports or audit-log entries); the Orchestrator decides what to surface."),
    ("D-033", "Velocity report and routing decisions are generated by deterministic scripts",
     "Aggregating numeric counts across JSON summaries is mechanical. Doing it in a LLM prompt burns tokens for arithmetic.",
     "build-velocity-report.ps1 produces the velocity report from manifest + audit-log + JSON summaries. route-defects.ps1 produces T-007 routing plan. Neither uses an LLM."),
    ("D-034", "Test sign-off currency enforced by hook + Orchestrator, not by the test agent",
     "Deciding WHEN to re-test is a coordination concern, not a tester-skill concern. Putting it in the test agent's prompt re-derives logic that belongs in deterministic code.",
     "When test agent reports PASS, Orchestrator invokes hook with -CommitSignoff which writes .signoff-hash. After every developer completion, Orchestrator re-invokes the test hook in re-execution mode: NO_CHANGE (sign-off current) or PROCEED (re-test)."),
    ("D-035", "Dispute authoring contract is producer-side documented",
     "If the dispute mechanism lives only in core KB and orchestrator docs, the agent that AUTHORS the dispute has no contract in its own files. The producer-side contract must be where the artefact is written.",
     "Add a SKILL — Defect Dispute Authoring section to each developer agent's skills file. Cover decision logic (fix vs dispute), frontmatter schema, JSON summary shape, re-judgement handshake, escalation paths."),
    ("D-036", "Post-completion validators invoked via -PostCheck switch on each producer hook",
     "Without an explicit invocation mechanism, post-completion validators are 'documented but not wired'. The Orchestrator needs one consistent call shape that emits one signal.",
     "Every producer hook accepts -PostCheck. Orchestrator invokes with this switch immediately after the agent reports complete and BEFORE marking the task [x]. Hook runs declared validators and emits VALIDATION_PASS / VALIDATION_FAIL."),
    ("D-037", "Manifest log triggers are documented in the Orchestrator definition",
     "Helper scripts (manifest-writer.ps1) exist but trigger events were implicit. Result: manifest sections stay empty even when the underlying events fire and the manifest stops being a live coordination record.",
     "Add a 'Manifest log triggers' subsection to the A-00 definition mapping pipeline events to helper functions (Append-AuditLog, Append-TestDefect, Append-Dispute, Append-Validation, Append-CrossSprint)."),
]

NEW_PIPELINE_CHECKLIST = [
    "Define the four pillars in your own words. Confirm the project actually benefits from agentic delivery (see 'When to use').",
    "Map the delivery flow into discrete phases with named inputs and outputs at each boundary.",
    "List every specialist responsibility. Apply the single-responsibility test: if any item contains 'and', split it.",
    "Draft the Orchestrator's role first. Confirm it produces no business output.",
    "Decide the workspace layout — at minimum, a separation between long-lived code, sprint-scoped artefacts, and pipeline infrastructure.",
    "Pick a naming convention for agents, hooks, validators, and tasks. Document it before writing any.",
    "Write each agent's definition file (responsibility, input, output, escalation, signing).",
    "Write each agent's skills file (output format, quality standards, DoD checklist, worked examples).",
    "Write each agent's hook script (dependency checks, path resolution, hash, partial-output recovery, return signals).",
    "Identify schema-validatable outputs. Write a per-producer validator script for each.",
    "Identify joint contracts between agents. Write cross-cutting alignment validators for each.",
    "Define producer JSON summaries for any output the Orchestrator routes on (reviews, defects, disputes, cross-sprint refs).",
    "Define the sign-off gate composition. Every downstream consumer of requirement cards is a signing agent.",
    "Document the clarification escalation chain per agent — primary target, fallback, eventual human.",
    "Write the Orchestrator definition: manifest sections, task registry, sign-off gate, clarification routing, blocker management, sign-off currency, post-completion validation, manifest log triggers.",
    "Document every design decision with its rationale. Future maintainers must understand WHY, not just WHAT.",
    "Wire the cost discipline rules: hash-skip, persisted briefings, mode-switch default, sub-agent budget, /compact discipline.",
    "Declare a model tier per agent. State the dynamic-rule policy (typically one rule only).",
    "Stand up the manifest, audit log, NOTIFICATIONS, and helper scripts (manifest-writer, route-defects, velocity report, model selector).",
    "Run a dry sprint end-to-end on a small fixture. Confirm every hook returns the right signal, every validator is invoked, every manifest section gets populated.",
    "Document the resulting pipeline in a core KB. Mark anything project-specific as override-table candidates.",
]


# ---------------------------------------------------------------------------
# Doc builders
# ---------------------------------------------------------------------------

def build_handbook(out_path: Path) -> None:
    doc = Document()
    _configure_base_styles(doc)

    add_title_block(
        doc,
        "Agentic Delivery Pipeline",
        "Design Reference Handbook",
        "Architects designing or extending multi-agent software delivery pipelines.",
    )
    add_page_break(doc)

    # ============= PART I — Foundations =============
    add_h1(doc, "Part I — Foundations")

    add_h2(doc, "1. Why Agentic Delivery")
    add_para(doc, PHILOSOPHY_OPENING)
    add_para(doc, "")
    for heading, items in WHEN_TO_USE:
        add_h3(doc, heading)
        add_bullets(doc, items)

    add_h2(doc, "2. The Four Pillars")
    add_para(doc, "Four foundational beliefs the pipeline is built on. Every decision can be traced back to one or more pillars.")
    for name, body in FOUR_PILLARS:
        add_h3(doc, name)
        add_para(doc, body)

    add_h2(doc, "3. Twelve Guiding Principles")
    add_para(doc, "These are the operating rules that fall out of the four pillars. If a decision violates a principle, the principle wins.")
    add_bullets(doc, GUIDING_PRINCIPLES)
    add_page_break(doc)

    # ============= PART II — Reference Architecture =============
    add_h1(doc, "Part II — Reference Architecture")

    add_h2(doc, "4. The Reference Twelve-Agent Pipeline")
    add_para(doc, "A reference pipeline ships with twelve agents: two coordinators, eight specialists, and two reserved slots for additional test layers. Your project may add more, but the twelve below are the minimum a non-trivial delivery pipeline tends to need when end-to-end test coverage matters.")
    add_table_with_header(
        doc,
        ["ID", "Name", "Single Responsibility", "Signs RC?"],
        REFERENCE_AGENTS,
        col_widths_in=[0.6, 1.8, 4.0, 0.8],
    )
    add_para(doc, "Reserved IDs: A-09 (microservice tester) and A-10 (database tester) when the project extends below the API/BFF layer or requires schema testing.", italic=True)
    add_para(doc, "Sign-off gate composition: six signing agents (every downstream consumer of requirement cards). A-01r is not a signing agent (its role is resolution, not consumption). A-03a is not a signing agent (it does not consume RC cards).")

    add_h2(doc, "5. Agent Anatomy")
    add_para(doc, "Every specialist agent has the same six components. Missing any of them is a sign the agent is not pipeline-ready.")
    for name, body in AGENT_ANATOMY:
        add_h3(doc, name)
        add_para(doc, body)

    add_h2(doc, "6. The Three-Folder Workspace")
    add_para(doc, "Long-lived code, sprint-scoped artefacts, and pipeline infrastructure live in three distinct folders. Mixing them causes scope confusion, accidental overwrites, and broken cross-sprint isolation.")
    for name, body in THREE_FOLDER:
        add_h3(doc, name)
        add_para(doc, body)
    add_callout(doc, "Lazy-creation rule", LAZY_CREATION_RULE)

    add_h2(doc, "7. Artefact and Filename Conventions")
    add_para(doc, "Pipeline infrastructure files use prefixed, fully-named identifiers so a glance at any filename answers 'what is this and which agent owns it':")
    add_bullets(doc, [
        "A-<id>-<fullname>-definition.md — agent definition file",
        "A-<id>-<fullname>-skills.md — agent skills file",
        "H-<id>-<fullname>.ps1 — agent hook script",
        "V-<id>-<topic>.ps1 — per-producer validator",
        "V-shared-<topic>.ps1 — cross-cutting validator",
        "CLAUDE-A-<id>-<fullname>.md — activation file (the entry point)",
        "T-<id>-A-<id>-<fullname>-briefing.md — persisted briefing for one task activation",
    ])
    add_para(doc, "Sprint artefacts keep their content-type prefixes (RC, ED, CI, CR, DEF, TC, TR, DSP) and accumulate sequentially across the sprint:")
    add_bullets(doc, [
        "RC-### — requirement card (frontmatter + body)",
        "ED-### — endpoint design",
        "CI-### — component inventory entry",
        "CR-### — code-review finding (CR2-, CR3- on iterative rework)",
        "AR-### — architecture-review finding",
        "TC-<layer>-### — test case",
        "TR-<layer>-### — test result",
        "DEF-<layer>-### — defect (DEF-FE-2-, DEF-BFF-2- on iterative rework)",
        "DSP-<layer>-### — defect dispute",
        "CL-### — clarification request",
        "CNC-### — concern (source-material gap)",
        "HB-### — human blocker",
    ])
    add_page_break(doc)

    # ============= PART III — Operating Model =============
    add_h1(doc, "Part III — Operating Model")

    add_h2(doc, "8. The Five Universal Protocols")
    add_para(doc, "Every agent inherits five protocols. Protocols are non-negotiable operating rules — when an agent's behaviour conflicts with one of them, the protocol wins.")
    for name, body in PROTOCOLS:
        add_h3(doc, name)
        add_para(doc, body)

    add_h2(doc, "9. Coordination Patterns")
    add_para(doc, "Six coordination patterns that emerge across most pipeline instances. Each pattern has a name, a trigger, an actor sequence, and a stop condition.")
    for name, body in COORDINATION_PATTERNS:
        add_h3(doc, name)
        add_para(doc, body)

    add_h2(doc, "10. Cost Discipline")
    add_para(doc, "Every token spent on coordination is a token not spent on delivery. Operating at the lowest cost tier that achieves the outcome is a hard discipline — moving up a tier requires explicit justification.")
    add_h3(doc, "The cost hierarchy")
    add_table_with_header(
        doc,
        ["Tier", "Mechanism", "Relative cost"],
        COST_TIERS,
        col_widths_in=[0.6, 4.5, 1.2],
    )
    add_h3(doc, "When sub-agent spawn is permitted (the only exceptions)")
    for name, body in SPAWN_CASES:
        add_h3(doc, name)
        add_para(doc, body)
    add_h3(doc, "Anti-patterns — these are NOT valid spawn justifications")
    add_bullets(doc, SPAWN_ANTIPATTERNS)
    add_callout(
        doc,
        "Cost-reduction tactics when spawn is justified",
        "Pre-authorise Write/Edit before launching. Use direct-disk-write for large payloads. Emit ledger-first so output survives truncation. Pre-decide canonical source-of-truth for parallel rework. Use owner-tag routing so sub-agents filter their inbox. Run verification gates (lint + tests) BEFORE reporting complete.",
    )

    add_h2(doc, "11. Validation and Quality")
    add_para(doc, "Quality is enforced at three layers, each catching a different class of failure.")
    for name, body in VALIDATION_LAYERS:
        add_h3(doc, name)
        add_para(doc, body)
    add_callout(
        doc,
        "Hook return signals",
        "PROCEED — inputs changed, run the task. NO_CHANGE — inputs identical, exit immediately (no LLM cost). "
        "BLOCKED:<reason> — precondition failed. ALIGNMENT_CONFLICT — joint contract broken. "
        "VALIDATION_PASS — post-completion schema check OK. VALIDATION_FAIL — schema violation; agent re-activated with detail.",
    )
    add_page_break(doc)

    # ============= PART IV — Engineering Discipline =============
    add_h1(doc, "Part IV — Engineering Discipline")

    add_h2(doc, "12. State Lives in Files")
    add_para(doc, "Agent sessions are stateless. Anything you need to survive a session restart, a /compact, a sub-agent spawn, or a fresh-session reload MUST live in a file. Holding state in agent memory across sessions multiplies cost by the re-load factor at every mode switch.")
    add_bullets(doc, [
        "orchestrator-manifest.md — single source of truth for all task state, gate status, blockers, sign-offs.",
        "briefings/T-###-A-##-<fullname>-briefing.md — per-activation context handoff.",
        "audit-log.md — append-only event history.",
        "<output-folder>/.input-hash — idempotency marker.",
        "<output-folder>/.signoff-hash — test sign-off currency record.",
        "concerns/resolutions/CL-###-resolution.md — clarification audit trail.",
    ])

    add_h2(doc, "13. The Hook Architecture")
    add_para(doc, "Hooks are deterministic scripts (typically PowerShell) that run before and after every agent activation. They cost zero LLM tokens and absorb mechanical work that would otherwise burn into the agent's prompt budget.")
    add_h3(doc, "Pre-activation responsibilities")
    add_bullets(doc, [
        "Verify dependencies are complete (read manifest, match task status patterns).",
        "Resolve all input and output paths.",
        "Load secrets from environment or vault.",
        "Compute input hash; compare to stored hash; emit NO_CHANGE if unchanged.",
        "Clean partial output if hash file is missing (interrupted prior run).",
        "Invoke pre-activation alignment validators if the agent consumes joint contracts.",
        "Return a single signal: PROCEED, NO_CHANGE, BLOCKED, or ALIGNMENT_CONFLICT.",
    ])
    add_h3(doc, "Post-completion responsibilities (-PostCheck switch)")
    add_bullets(doc, [
        "Invoke declared Tier-1 validators against the agent's emitted artefacts.",
        "Collapse multiple validator exits into one signal: VALIDATION_PASS or VALIDATION_FAIL.",
        "The Orchestrator calls this BEFORE marking the task [x]; on FAIL, re-activate with the failure detail.",
    ])

    add_h2(doc, "14. Mode-Switch vs Sub-Agent Spawn")
    add_para(doc, "The pipeline distinguishes two activation mechanisms with very different cost profiles.")
    add_h3(doc, "Foreground mode-switch (default, Tier 1 cost)")
    add_para(doc, "When activated, the receiving session BECOMES the agent by reading its definition + skills + briefing. Same session, same model. State persists in files on disk between mode-switches. This is the cheapest mechanism and handles ~85% of activations.")
    add_h3(doc, "Sub-agent spawn (exception, Tier 4 cost, 3–5× baseline)")
    add_para(doc, "A separate sub-agent in isolated context. Permitted only for Case A/B/C (Section 10). Every spawn is logged with its Case justification. Default budget: 2 spawns per sprint. Exceeding the budget requires a documented Protocol 5 exception.")

    add_h2(doc, "15. Idempotency via Hash-Skip")
    add_para(doc, "The pipeline is safe to re-trigger at any point. Re-running on unchanged inputs produces zero output and zero cost.")
    add_bullets(doc, [
        "Each agent's hook computes MD5 of all relevant inputs.",
        "The combined hash is stored in <output-folder>/.input-hash (or sprint-scoped variant for app/ folders).",
        "On the next run, the hook compares current hash to stored hash. Match → NO_CHANGE → exit.",
        "Agents trust NO_CHANGE absolutely. They do not 'double-check' by re-running.",
    ])

    add_h2(doc, "16. Producer JSON Summaries for Routing")
    add_para(doc, "The Orchestrator reads JSON summaries for every routing decision. It NEVER parses Markdown artefact bodies. This decouples coordination logic from artefact format and keeps routing fast and deterministic.")
    add_bullets(doc, [
        "review-summary.json — code-review verdict + finding counts + rework-required flag.",
        "defect-summary-<layer>.json — defects by owner/criticality + coverage + verdict + rework flag.",
        "dispute-summary.json — disputes by defectRef + verdict + disputer.",
        "cross-sprint-refs.json — requirement-card references that touch other sprints.",
        "T-007-routing-plan.json — Orchestrator-emitted unified rework plan.",
    ])

    add_h2(doc, "17. Mechanical Helpers")
    add_para(doc, "Reusable scripts that hooks and the Orchestrator invoke. They hold mechanical detail so agent definitions stay narrative.")
    add_table_with_header(
        doc,
        ["Script", "Purpose"],
        [
            ("start-sprint", "Bootstrap a new sprint (creates sprint folder + START_SPRINT signal)."),
            ("workspace-setup", "First-time workspace bootstrap."),
            ("setup-secrets", "First-time credential walkthrough."),
            ("manifest-writer", "Helpers to append rows to manifest tables (Audit Log, Test Defect Log, Dispute Log, Validation Log, Cross-Sprint Log)."),
            ("route-defects", "Reads review-summary + defect-summary JSONs, emits unified T-007 routing plan."),
            ("build-velocity-report", "Generates sprint velocity report from manifest + audit-log + JSON summaries."),
            ("build-review-report", "Converts rework ledger JSON to Excel for human review."),
            ("select-model", "Picks model tier for a sub-agent spawn (declared tier + one dynamic rule)."),
            ("check-scripts", "Parse-checks all scripts in the pipeline."),
        ],
        col_widths_in=[1.8, 4.7],
    )
    add_page_break(doc)

    # ============= PART V — Failure Modes & Recovery =============
    add_h1(doc, "Part V — Failure Modes and Recovery")

    add_h2(doc, "18. Common Failure Modes")
    add_para(doc, "Each failure mode is captured as Symptom / Cause / Prevention / Detection / Recovery. Use this catalogue when designing your pipeline's guardrails and when triaging real incidents.")
    for name, sym, prevent, recov in FAILURE_MODES:
        add_h3(doc, name)
        add_para(doc, sym)
        p = doc.add_paragraph()
        r = p.add_run("Prevention / Detection: ")
        r.bold = True
        p.add_run(prevent)
        p = doc.add_paragraph()
        r = p.add_run("Recovery: ")
        r.bold = True
        p.add_run(recov)

    add_h2(doc, "19. Truncation Recovery")
    add_para(doc, "When a sub-agent's return exceeds the transcript ceiling, the tail is silently truncated. Three tactics prevent and recover from this:")
    add_bullets(doc, [
        "Ledger-first emission — the per-agent ledger JSON is the FIRST artefact in the response, not the last. Even if the body is truncated, the routing contract survives.",
        "Direct-disk-write — pre-auth Write/Edit in settings.local.json, then the sub-agent writes to disk and returns only a summary. The persisted file is what matters, not the transcript.",
        "Focused split — instead of one wide-scope sub-agent, fire 2–3 narrow-scope sub-agents in parallel. Each fits under the ceiling.",
    ])

    add_h2(doc, "20. Anti-patterns Catalogue")
    add_para(doc, "Patterns that look reasonable but actively damage pipeline reliability or cost. Each is treated as a defect when observed.")
    add_bullets(doc, ANTIPATTERNS)
    add_page_break(doc)

    # ============= PART VI — Scaling & Adaptation =============
    add_h1(doc, "Part VI — Scaling and Adaptation")

    add_h2(doc, "21. Adding a New Specialist Agent")
    add_para(doc, "When the pipeline needs a new responsibility — a security reviewer, a performance tester, an API spec writer — follow this thirteen-step pattern.")
    add_numbered(doc, [
        "Define the agent's single responsibility — answer the three SRP questions.",
        "Pick the next sequential <id>, or a suffix like '01r' for a sibling split.",
        "Pick a <fullname> (lowercase, hyphenated, descriptive, 1–3 words).",
        "Identify the agent's position in the dependency graph.",
        "Create the definition file with all six required components.",
        "Create the skills file with output format, quality standards, DoD checklist.",
        "Create the hook script covering all standard pre-activation and (if applicable) post-completion responsibilities.",
        "Create the activation file with a Default model tier section.",
        "Update the Orchestrator manifest: FOLDER REGISTRY, CLARIFICATION ESCALATION CHAIN, TASK REGISTRY.",
        "If the new agent consumes requirement cards, add it to the sign-off gate.",
        "If it emits routable output, define the JSON summary contract.",
        "If it has a schema-validatable output, create a V-<id>-<topic>.ps1 validator.",
        "Update the cost-discipline KB's model-tier table.",
    ])

    add_h2(doc, "22. Reducing to a Minimal Pipeline")
    add_para(doc, "For very small projects, the twelve-agent shape is overkill. A four-agent minimum: Orchestrator + Requirement Analyst (no split) + one Developer + one Reviewer. Sign-off gate simplified to one signing agent. No sprint structure required. The five universal protocols still apply — reduction is in agent count, not in discipline.")

    add_h2(doc, "23. Multiple Developers in Parallel")
    add_para(doc, "When a single developer agent becomes a bottleneck:")
    add_bullets(doc, [
        "Create one dedicated hooks script per developer persona (e.g. H-04a, H-04b).",
        "Split the requirement card range across personas (A-04a owns RC-011..015; A-04b owns RC-016..020).",
        "The Orchestrator activates both in parallel after the gate opens.",
        "Each developer writes to its own subdirectory within the sprint output folder.",
        "The Code Reviewer receives all output combined.",
    ])

    add_h2(doc, "24. Adding New Test Layers")
    add_para(doc, "When the application extends below the API/BFF layer (microservices, databases, integration partners):")
    add_bullets(doc, [
        "Reserve A-09 for microservice tester, A-10 for database tester.",
        "Follow the test-agent dual-phase pattern (planning + execution).",
        "Use the tests/<layer>/{test-cases,test-results/{defects,disputes}}/ folder structure.",
        "Emit defect-summary-<layer>.json per the producer JSON summary contract.",
        "Add the new test agent to the sign-off gate.",
        "Update T-007 routing-plan logic to include the new layer's defects.",
    ])
    add_page_break(doc)

    # ============= PART VII — Reference =============
    add_h1(doc, "Part VII — Reference")

    add_h2(doc, "25. Decision Log")
    add_para(doc, "Every load-bearing design decision in the reference pipeline, with its rationale and practical application. New pipelines should produce a similar log of their own decisions — undocumented decisions become mysteries.")
    for did, title, why, how_to in DECISIONS:
        add_decision(doc, did, title, why, how_to)

    add_h2(doc, "26. Glossary")
    add_table_with_header(
        doc,
        ["Term", "Definition"],
        [
            ("Agent", "A specialist AI process with a single defined responsibility, its own definition + skills + hook + activation files, and a persisted briefing per activation."),
            ("Briefing", "Per-task context handoff persisted to disk before agent activation. The durable record."),
            ("DoD", "Definition of Done. The checklist in an agent's skills file that the agent must self-validate before reporting complete."),
            ("Foreground mode-switch", "Activation where the receiving session adopts the agent role in the same session. Tier-1 cost."),
            ("Gate", "The requirement sign-off gate. All signing agents review RC cards READ-ONLY; opens when all sign-offs are collected."),
            ("Hash-skip", "Hook returns NO_CHANGE when input hash matches the stored hash. Zero-LLM-cost no-op."),
            ("Hook", "Deterministic script (typically PowerShell) wrapping every agent activation. Pre-activation checks + optional post-completion validation."),
            ("Manifest", "orchestrator-manifest.md. The single source of truth for all pipeline state."),
            ("Mode-switch", "See Foreground mode-switch."),
            ("Orchestrator", "A-00. The central coordinator. Single writer of the manifest. Produces no business output."),
            ("Protocol", "One of the five universal operating rules every agent inherits."),
            ("Skill file", "An agent's HOW. Domain knowledge, output spec, quality standards, DoD."),
            ("Sub-agent spawn", "Activation in an isolated context. Tier-4 cost. Permitted only for Case A/B/C."),
            ("Validator", "Deterministic script that mechanically checks a schema or joint contract. Per-producer or cross-cutting."),
        ],
        col_widths_in=[1.5, 5.0],
    )

    add_h2(doc, "27. Checklist — Designing a New Pipeline")
    add_para(doc, "Use this checklist as a one-pass scaffolding tool when starting a new project. Each step has a corresponding section earlier in the handbook for depth.")
    add_numbered(doc, NEW_PIPELINE_CHECKLIST)

    doc.save(str(out_path))
    print(f"  wrote {out_path}")


def build_playbook(out_path: Path) -> None:
    doc = Document()
    _configure_base_styles(doc)

    add_title_block(
        doc,
        "Agentic Delivery Pipeline",
        "Concise Design Playbook",
        "Project leads and architects who need the load-bearing rules without the full handbook.",
    )
    add_page_break(doc)

    add_h1(doc, "1. Philosophy in One Page")
    add_para(doc, PHILOSOPHY_OPENING)
    for name, body in FOUR_PILLARS:
        add_h2(doc, name)
        add_para(doc, body)

    add_h1(doc, "2. Twelve Guiding Principles")
    add_bullets(doc, GUIDING_PRINCIPLES)

    add_h1(doc, "3. Reference Architecture")
    add_para(doc, "Twelve agents — two coordinators, eight specialists, two reserved. The shape, not the law: smaller projects collapse to four agents; larger projects fan out parallel developer personas.")
    add_table_with_header(
        doc,
        ["ID", "Name", "Responsibility", "Signs RC?"],
        REFERENCE_AGENTS,
        col_widths_in=[0.6, 1.8, 4.0, 0.8],
    )
    add_callout(doc, "Lazy-creation rule (workspace layout)", LAZY_CREATION_RULE)

    add_h1(doc, "4. The Five Universal Protocols")
    for name, body in PROTOCOLS:
        add_h2(doc, name)
        add_para(doc, body)

    add_h1(doc, "5. Cost Discipline at a Glance")
    add_h2(doc, "Cost hierarchy")
    add_table_with_header(
        doc,
        ["Tier", "Mechanism", "Relative cost"],
        COST_TIERS,
        col_widths_in=[0.6, 4.5, 1.2],
    )
    add_h2(doc, "When sub-agent spawn is permitted")
    for name, body in SPAWN_CASES:
        add_h3(doc, name)
        add_para(doc, body)

    add_h1(doc, "6. Curated Load-Bearing Decisions")
    add_para(doc, "The decisions whose rationale tends to be re-discovered painfully when ignored. The full decision log lives in the Design Handbook.")
    load_bearing = [
        "D-001", "D-002", "D-004", "D-006", "D-009", "D-010", "D-011",
        "D-019", "D-022", "D-023", "D-024", "D-025", "D-028", "D-029",
        "D-030", "D-034", "D-035", "D-036", "D-037",
    ]
    by_id = {d[0]: d for d in DECISIONS}
    for did in load_bearing:
        d = by_id[did]
        add_decision(doc, d[0], d[1], d[2], d[3])

    add_h1(doc, "7. Top Anti-patterns to Watch For")
    add_bullets(doc, ANTIPATTERNS[:12])

    add_h1(doc, "8. New Pipeline Design Checklist")
    add_numbered(doc, NEW_PIPELINE_CHECKLIST)

    doc.save(str(out_path))
    print(f"  wrote {out_path}")


def build_cheatsheet(out_path: Path) -> None:
    doc = Document()
    _configure_base_styles(doc)

    # Tighter margins for the cheat-sheet
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Agentic Delivery Pipeline — Cheat-Sheet")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = COLOR_PRIMARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"v{VERSION} · {TODAY} · project-independent quick reference")
    r.italic = True
    r.font.color.rgb = COLOR_MUTED

    # Workspace layout (lazy-creation rule)
    add_h2(doc, "Workspace Layout (Lazy-Creation Rule)")
    add_table_with_header(
        doc,
        ["Folder", "Present at bootstrap?", "Created by"],
        [
            ("agentic-pipeline/", "YES — only folder in a fresh workspace", "workspace-setup.ps1"),
            ("sprints/", "no — appears on first sprint", "start-sprint.ps1"),
            ("app/frontend/", "no — appears on first FE activation", "H-04 hook"),
            ("app/backend/", "no — appears on first BFF activation", "H-05 hook"),
        ],
        col_widths_in=[1.8, 2.7, 2.0],
    )

    # Four Pillars (compact)
    add_h2(doc, "Four Pillars")
    add_table_with_header(
        doc,
        ["Pillar", "One-line"],
        [
            ("Single Responsibility", "One agent, one job. 'and' in a responsibility means split it."),
            ("Centralised Coordination", "Orchestrator holds all state. Agents never contact each other directly."),
            ("Explicit over Implicit", "No guesses, no assumptions. Every ambiguity is a clarification."),
            ("Cost Discipline", "Cheapest tier that achieves the outcome. Mode-switch over spawn."),
        ],
        col_widths_in=[1.8, 4.7],
    )

    # Five Protocols (compact)
    add_h2(doc, "Five Universal Protocols")
    add_table_with_header(
        doc,
        ["#", "Protocol", "One-line"],
        [
            ("1", "Startup", "Read briefing. Confirm input/output paths + dependency status. Trust NO_CHANGE."),
            ("2", "Sign-off", "Signing agents review RC cards READ-ONLY at the gate before any design begins."),
            ("3", "Clarification", "Never guess. Raise CL via Orchestrator; route per escalation chain."),
            ("4", "Completion", "Self-validate DoD. Hook -PostCheck validates schemas before manifest marks [x]."),
            ("5", "Cost Discipline", "Mode-switch default. Hash-skip absolutely. Sub-agent spawn only for Case A/B/C."),
        ],
        col_widths_in=[0.3, 1.5, 4.7],
    )

    # Cost hierarchy
    add_h2(doc, "Cost Hierarchy")
    add_table_with_header(
        doc,
        ["Tier", "Mechanism", "Cost"],
        COST_TIERS,
        col_widths_in=[0.5, 4.8, 1.2],
    )

    # When to spawn
    add_h2(doc, "Sub-agent Spawn — The Only Three Cases")
    add_table_with_header(
        doc,
        ["Case", "Trigger"],
        [
            ("A — True parallelism", "Two agents must run simultaneously; sequencing would idle one."),
            ("B — Heavy context isolation", "Exploration would blow up the foreground context (e.g. 100+ file scan)."),
            ("C — Truncation-risk payload", "Single emission would exceed the transcript ceiling."),
        ],
        col_widths_in=[1.8, 4.7],
    )

    # Hook signals
    add_h2(doc, "Hook Return Signals")
    add_table_with_header(
        doc,
        ["Signal", "Meaning"],
        [
            ("PROCEED", "Inputs changed (or first run). Activate the agent."),
            ("NO_CHANGE", "Inputs identical. Exit; mark task [=] Skipped."),
            ("BLOCKED:<reason>", "Precondition failed (gate not open, missing input)."),
            ("ALIGNMENT_CONFLICT", "Pre-activation joint-contract check failed. Route back upstream."),
            ("VALIDATION_PASS", "Post-completion validators passed. Safe to mark [x]."),
            ("VALIDATION_FAIL", "Schema violation. Re-activate with failure detail in next briefing."),
        ],
        col_widths_in=[1.8, 4.7],
    )

    # Anti-patterns top 10
    add_h2(doc, "Top Anti-patterns")
    add_bullets(doc, ANTIPATTERNS[:10])

    # New pipeline 13-step checklist (compact form)
    add_h2(doc, "New Pipeline — 13-Step Scaffold")
    add_numbered(doc, [
        "Confirm agentic delivery fits — separable phases, quality gates matter, multiple specialists.",
        "Apply the SRP test to every proposed agent. Split anything with 'and'.",
        "Decide the workspace layout (long-lived code / sprint artefacts / pipeline infra).",
        "Pick the naming convention. Document it before writing files.",
        "Write definition + skills + hook + activation per agent.",
        "Write per-producer + cross-cutting validators.",
        "Define producer JSON summaries for routable outputs.",
        "Define sign-off gate composition and clarification escalation chains.",
        "Write the Orchestrator definition; wire post-completion -PostCheck and manifest log triggers.",
        "Wire cost discipline: hash-skip, briefings, mode-switch default, sub-agent budget, /compact.",
        "Declare model tier per agent; state the dynamic-rule policy.",
        "Stand up manifest, audit log, NOTIFICATIONS, helper scripts.",
        "Dry-run a sprint on a small fixture; confirm every hook + validator fires correctly.",
    ])

    doc.save(str(out_path))
    print(f"  wrote {out_path}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main() -> None:
    script_dir = Path(__file__).resolve().parent          # …/agentic-pipeline/scripts
    workspace_root = script_dir.parent.parent             # …/poc-workspace

    out_handbook = workspace_root / "Agentic-Pipeline-Design-Handbook.docx"
    out_playbook = workspace_root / "Agentic-Pipeline-Concise-Playbook.docx"
    out_cheat = workspace_root / "Agentic-Pipeline-Cheat-Sheet.docx"

    print("Building three project-independent design reference docs:")
    build_handbook(out_handbook)
    build_playbook(out_playbook)
    build_cheatsheet(out_cheat)
    print("Done.")


if __name__ == "__main__":
    main()
